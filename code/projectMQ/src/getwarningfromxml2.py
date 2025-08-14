import xml.etree.ElementTree as ET
from docx import Document

def parse_xml(file_path):
    tree = ET.parse(file_path)
    root = tree.getroot()
    errors = []

    for error in root.findall('Warning'):
        error_data = {}
        for attribute in error.findall('Attribute'):
            name = attribute.get('name')
            value = attribute.get('value')
            if name in ['warning_code', 'warning_msg', 'warning_solution']:
                error_data[name] = value
        errors.append(error_data)
    return errors

def create_word_document(errors, output_path):
    doc = Document()
    for error in errors:
        doc.add_heading(f"Error Code: {error['warning_code']}", level=1)
        doc.add_paragraph(f"Description: {error['warning_msg']}")
        doc.add_paragraph(f"Solution: {error['warning_solution']}")
    doc.save(output_path)

if __name__ == "__main__":
    xml_file = '/opt/wupei/git/code/projectMQ/src/warningCode-0221.xml'
    output_docx = '/opt/wupei/git/code/projectMQ/src/warningCode-0221.docx'
    
    errors = parse_xml(xml_file)
    create_word_document(errors, output_docx)
    print(f"Content has been imported into {output_docx}")