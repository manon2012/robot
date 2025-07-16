# #coding=utf-8
# from xml.etree import ElementTree as ET
# per=ET.parse('errorcode.xml')
# p=per.findall('./login/item')

# for oneper in p:
#     for child in oneper.getchildren():
#         #print child.tag,':',child.text
#         print(child)


# # p=per.findall('./item')

# # for oneper in p:
# #     for child in oneper.getchildren():
# #         print child.tag,':',child.text

# import xml.etree.ElementTree as ET
# from docx import Document

# def parse_xml(file_path):
#     tree = ET.parse(file_path)
#     root = tree.getroot()
#     items = []
#     for item in root.findall('Error'):
#         title = item.find('error_code').text
#         description = item.find('error_description').text
#         items.append((title, description))
#     return items

# def create_word_document(items, output_path):
#     doc = Document()
#     for title, description in items:
#         doc.add_heading(title, level=1)
#         doc.add_paragraph(description)
#     doc.save(output_path)

# if __name__ == "__main__":
#     xml_file = '/opt/wupei/git/code/projectMQ/src/errorcode.xml'
#     output_docx = '/opt/wupei/git/code/projectMQ/src/output.docx'
    
#     items = parse_xml(xml_file)
#     create_word_document(items, output_docx)
#     print(f"Content has been imported into {output_docx}")




import xml.etree.ElementTree as ET
from docx import Document

def parse_xml(file_path):
    tree = ET.parse(file_path)
    root = tree.getroot()
    items = []
    for item in root.findall('item'):
        title = item.find('title').text
        description = item.find('description').text
        items.append((title, description))
    return items

def create_word_document(items, output_path):
    doc = Document()
    for title, description in items:
        doc.add_heading(title, level=1)
        doc.add_paragraph(description)
    doc.save(output_path)

if __name__ == "__main__":
    xml_file = '/opt/wupei/git/code/projectMQ/src/sample.xml'
    output_docx = '/opt/wupei/git/code/projectMQ/src/output.docx'
    
    items = parse_xml(xml_file)
    create_word_document(items, output_docx)
    print(f"Content has been imported into {output_docx}")